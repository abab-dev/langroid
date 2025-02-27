from __future__ import annotations

import itertools
import logging
import os
import re
import tempfile
from enum import Enum
from io import BytesIO
from itertools import accumulate
from pathlib import Path
from typing import TYPE_CHECKING, Any, Dict, Generator, List, Tuple

from langroid.exceptions import LangroidImportError
from langroid.utils.object_registry import ObjectRegistry

try:
    import fitz
except ImportError:
    if not TYPE_CHECKING:
        fitz = None
try:
    import pymupdf4llm
except ImportError:
    if not TYPE_CHECKING:
        pymupdf4llm = None

try:
    import docling
except ImportError:
    if not TYPE_CHECKING:
        docling = None

try:
    import pypdf
except ImportError:
    if not TYPE_CHECKING:
        pypdf = None


import requests
from bs4 import BeautifulSoup

if TYPE_CHECKING:
    from PIL import Image

from langroid.mytypes import DocMetaData, Document
from langroid.parsing.parser import Parser, ParsingConfig

logger = logging.getLogger(__name__)


class DocumentType(str, Enum):
    # TODO add `md` (Markdown) and `html`
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    TXT = "txt"
    XLSX = "xlsx"
    XLS = "xls"
    PPTX = "pptx"


def find_last_full_char(possible_unicode: bytes) -> int:
    """
    Find the index of the last full character in a byte string.
    Args:
        possible_unicode (bytes): The bytes to check.
    Returns:
        int: The index of the last full unicode character.
    """

    for i in range(len(possible_unicode) - 1, 0, -1):
        if (possible_unicode[i] & 0xC0) != 0x80:
            return i
    return 0


def is_plain_text(path_or_bytes: str | bytes) -> bool:
    """
    Check if a file is plain text by attempting to decode it as UTF-8.
    Args:
        path_or_bytes (str|bytes): The file path or bytes object.
    Returns:
        bool: True if the file is plain text, False otherwise.
    """
    if isinstance(path_or_bytes, str):
        if path_or_bytes.startswith(("http://", "https://")):
            response = requests.get(path_or_bytes)
            response.raise_for_status()
            content = response.content[:1024]
        else:
            with open(path_or_bytes, "rb") as f:
                content = f.read(1024)
    else:
        content = path_or_bytes[:1024]
    try:
        # Use magic to detect the MIME type
        import magic

        mime_type = magic.from_buffer(content, mime=True)

        # Check if the MIME type is not a text type
        if not mime_type.startswith("text/"):
            return False

        # Attempt to decode the content as UTF-8
        content = content[: find_last_full_char(content)]

        try:
            _ = content.decode("utf-8")
            # Additional checks can go here, e.g., to verify that the content
            # doesn't contain too many unusual characters for it to be considered text
            return True
        except UnicodeDecodeError:
            return False
    except UnicodeDecodeError:
        # If decoding fails, it's likely not plain text (or not encoded in UTF-8)
        return False


class DocumentParser(Parser):
    """
    Abstract base class for extracting text from special types of docs
    such as PDFs or Docx.

    Attributes:
        source (str): The source, either a URL or a file path.
        doc_bytes (BytesIO): BytesIO object containing the doc data.
    """

    @classmethod
    def create(
        cls,
        source: str | bytes,
        config: ParsingConfig,
        doc_type: str | DocumentType | None = None,
    ) -> "DocumentParser":
        """
        Create a DocumentParser instance based on source type
            and config.<source_type>.library specified.

        Args:
            source (str|bytes): The source, could be a URL, file path,
                or bytes object.
            config (ParserConfig): The parser configuration.
            doc_type (str|None): The type of document, if known

        Returns:
            DocumentParser: An instance of a DocumentParser subclass.
        """
        inferred_doc_type = DocumentParser._document_type(source, doc_type)
        if inferred_doc_type == DocumentType.PDF:
            if config.pdf.library == "fitz":
                return FitzPDFParser(source, config)
            elif config.pdf.library == "pymupdf4llm":
                return PyMuPDF4LLMParser(source, config)
            elif config.pdf.library == "docling":
                return DoclingParser(source, config)
            elif config.pdf.library == "pypdf":
                return PyPDFParser(source, config)
            elif config.pdf.library == "unstructured":
                return UnstructuredPDFParser(source, config)
            elif config.pdf.library == "pdf2image":
                return ImagePdfParser(source, config)
            else:
                raise ValueError(
                    f"Unsupported PDF library specified: {config.pdf.library}"
                )
        elif inferred_doc_type == DocumentType.DOCX:
            if config.docx.library == "unstructured":
                return UnstructuredDocxParser(source, config)
            elif config.docx.library == "python-docx":
                return PythonDocxParser(source, config)
            else:
                raise ValueError(
                    f"Unsupported DOCX library specified: {config.docx.library}"
                )
        elif inferred_doc_type == DocumentType.DOC:
            return UnstructuredDocParser(source, config)
        elif inferred_doc_type == DocumentType.XLS:
            return MarkitdownXLSXParser(source, config)
        elif inferred_doc_type == DocumentType.XLSX:
            return MarkitdownXLSXParser(source, config)
        elif inferred_doc_type == DocumentType.PPTX:
            return MarkitdownPPTXParser(source, config)
        else:
            source_name = source if isinstance(source, str) else "bytes"
            raise ValueError(f"Unsupported document type: {source_name}")

    def __init__(self, source: str | bytes, config: ParsingConfig):
        """
        Args:
            source (str|bytes): The source, which could be
            a path, a URL or a bytes object.
        """
        super().__init__(config)
        self.config = config
        if isinstance(source, bytes):
            self.source = "bytes"
            self.doc_bytes = BytesIO(source)
        else:
            self.source = source
            self.doc_bytes = self._load_doc_as_bytesio()

    @staticmethod
    def _document_type(
        source: str | bytes, doc_type: str | DocumentType | None = None
    ) -> DocumentType:
        """
        Determine the type of document based on the source.

        Args:
            source (str|bytes): The source, which could be a URL,
                a file path, or a bytes object.
            doc_type (str|DocumentType|None): The type of document, if known.

        Returns:
            str: The document type.
        """
        if isinstance(doc_type, DocumentType):
            return doc_type
        if doc_type:
            return DocumentType(doc_type.lower())
        if is_plain_text(source):
            return DocumentType.TXT
        if isinstance(source, str):
            # detect file type from path extension
            if source.lower().endswith(".pdf"):
                return DocumentType.PDF
            elif source.lower().endswith(".docx"):
                return DocumentType.DOCX
            elif source.lower().endswith(".doc"):
                return DocumentType.DOC
            elif source.lower().endswith(".xlsx"):
                return DocumentType.XLSX
            elif source.lower().endswith(".xls"):
                return DocumentType.XLS
            elif source.lower().endswith(".pptx"):
                return DocumentType.PPTX
            else:
                raise ValueError(f"Unsupported document type: {source}")
        else:
            # must be bytes: attempt to detect type from content
            # using magic mime type detection
            import magic

            mime_type = magic.from_buffer(source, mime=True)
            if mime_type == "application/pdf":
                return DocumentType.PDF
            elif mime_type in [
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document",
            ]:
                return DocumentType.DOCX
            elif mime_type == "application/msword":
                return DocumentType.DOC
            elif (
                mime_type
                == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            ):
                return DocumentType.XLSX
            elif mime_type == "application/vnd.ms-excel":
                return DocumentType.XLS
            else:
                raise ValueError("Unsupported document type from bytes")

    def _load_doc_as_bytesio(self) -> BytesIO:
        """
        Load the docs into a BytesIO object.

        Returns:
            BytesIO: A BytesIO object containing the doc data.
        """
        if self.source.startswith(("http://", "https://")):
            response = requests.get(self.source)
            response.raise_for_status()
            return BytesIO(response.content)
        else:
            with open(self.source, "rb") as f:
                return BytesIO(f.read())

    @staticmethod
    def chunks_from_path_or_bytes(
        source: str | bytes,
        parser: Parser,
        doc_type: str | DocumentType | None = None,
        lines: int | None = None,
    ) -> List[Document]:
        """
        Get document chunks from a file path or bytes object.
        Args:
            source (str|bytes): The source, which could be a URL, path or bytes object.
            parser (Parser): The parser instance (for splitting the document).
            doc_type (str|DocumentType|None): The type of document, if known.
            lines (int|None): The number of lines to read from a plain text file.
        Returns:
            List[Document]: A list of `Document` objects,
                each containing a chunk of text, determined by the
                chunking and splitting settings in the parser config.
        """
        dtype: DocumentType = DocumentParser._document_type(source, doc_type)
        if dtype in [
            DocumentType.PDF,
            DocumentType.DOC,
            DocumentType.DOCX,
            DocumentType.PPTX,
            DocumentType.XLS,
            DocumentType.XLSX,
        ]:
            doc_parser = DocumentParser.create(
                source,
                parser.config,
                doc_type=doc_type,
            )
            chunks = doc_parser.get_doc_chunks()
            if len(chunks) == 0 and dtype == DocumentType.PDF:
                doc_parser = ImagePdfParser(source, parser.config)
                chunks = doc_parser.get_doc_chunks()
            return chunks
        else:
            # try getting as plain text; these will be chunked downstream
            # -- could be a bytes object or a path
            if isinstance(source, bytes):
                content = source.decode()
                if lines is not None:
                    file_lines = content.splitlines()[:lines]
                    content = "\n".join(line.strip() for line in file_lines)
            else:
                with open(source, "r") as f:
                    if lines is not None:
                        file_lines = list(itertools.islice(f, lines))
                        content = "\n".join(line.strip() for line in file_lines)
                    else:
                        content = f.read()
            soup = BeautifulSoup(content, "html.parser")
            text = soup.get_text()
            source_name = source if isinstance(source, str) else "bytes"
            doc = Document(
                content=text,
                metadata=DocMetaData(source=str(source_name)),
            )
            return parser.split([doc])

    def iterate_pages(self) -> Generator[Tuple[int, Any], None, None]:
        """Yield each page in the PDF."""
        raise NotImplementedError

    def get_document_from_page(self, page: Any) -> Document:
        """
        Get Langroid Document object (with possible metadata)
        corresponding to a given page.
        """
        raise NotImplementedError

    def fix_text(self, text: str) -> str:
        """
        Fix text extracted from a PDF.

        Args:
            text (str): The extracted text.

        Returns:
            str: The fixed text.
        """
        # Some pdf parsers introduce extra space before hyphen,
        # so use regular expression to replace 'space-hyphen' with just 'hyphen'
        return re.sub(r" +\-", "-", text)

    def get_doc(self) -> Document:
        """
        Get entire text from source as a single document.

        Returns:
            a `Document` object containing the content of the pdf file,
                and metadata containing source name (URL or path)
        """

        text = "".join(
            [
                self.get_document_from_page(page).content
                for _, page in self.iterate_pages()
            ]
        )
        return Document(content=text, metadata=DocMetaData(source=self.source))

    def get_doc_chunks(self) -> List[Document]:
        """
        Get document chunks from a pdf source,
        with page references in the document metadata.

        Adapted from
        https://github.com/whitead/paper-qa/blob/main/paperqa/readers.py

        Returns:
            List[Document]: a list of `Document` objects,
                each containing a chunk of text
        """

        split = []  # tokens in curr split
        pages: List[str] = []
        docs: List[Document] = []
        # metadata.id to be shared by ALL chunks of this document
        common_id = ObjectRegistry.new_id()
        n_chunks = 0  # how many chunk so far
        for i, page in self.iterate_pages():
            # not used but could be useful, esp to blend the
            # metadata from the pages into the chunks
            page_doc = self.get_document_from_page(page)
            page_text = page_doc.content
            split += self.tokenizer.encode(page_text)
            pages.append(str(i + 1))
            # split could be so long it needs to be split
            # into multiple chunks. Or it could be so short
            # that it needs to be combined with the next chunk.
            while len(split) > self.config.chunk_size:
                # pretty formatting of pages (e.g. 1-3, 4, 5-7)
                p_0 = int(pages[0])
                p_n = int(pages[-1])
                page_str = f"pages {p_0}-{p_n}" if p_0 != p_n else f"page {p_0}"
                text = self.tokenizer.decode(split[: self.config.chunk_size])
                docs.append(
                    Document(
                        content=text,
                        metadata=DocMetaData(
                            source=f"{self.source} {page_str}",
                            is_chunk=True,
                            id=common_id,
                        ),
                    )
                )
                n_chunks += 1
                split = split[self.config.chunk_size - self.config.overlap :]
                pages = [str(i + 1)]
        # there may be a last split remaining:
        # if it's shorter than the overlap, we shouldn't make a chunk for it
        # since it's already included in the prior chunk;
        # the only exception is if there have been no chunks so far.
        if len(split) > self.config.overlap or n_chunks == 0:
            pg = "-".join([pages[0], pages[-1]])
            text = self.tokenizer.decode(split[: self.config.chunk_size])
            docs.append(
                Document(
                    content=text,
                    metadata=DocMetaData(
                        source=f"{self.source} pages {pg}",
                        is_chunk=True,
                        id=common_id,
                    ),
                )
            )
        self.add_window_ids(docs)
        return docs


class FitzPDFParser(DocumentParser):
    """
    Parser for processing PDFs using the `fitz` library.
    """

    def iterate_pages(self) -> Generator[Tuple[int, "fitz.Page"], None, None]:
        """
        Yield each page in the PDF using `fitz`.

        Returns:
            Generator[fitz.Page]: Generator yielding each page.
        """
        if fitz is None:
            raise LangroidImportError("fitz", "pdf-parsers")
        doc = fitz.open(stream=self.doc_bytes, filetype="pdf")
        for i, page in enumerate(doc):
            yield i, page
        doc.close()

    def get_document_from_page(self, page: "fitz.Page") -> Document:
        """
        Get Document object from a given `fitz` page.

        Args:
            page (fitz.Page): The `fitz` page object.

        Returns:
            Document: Document object, with content and possible metadata.
        """
        return Document(
            content=self.fix_text(page.get_text()),
            metadata=DocMetaData(source=self.source),
        )


class PyMuPDF4LLMParser(DocumentParser):
    """
    Parser for processing PDFs using the `pymupdf4llm` library.
    """

    def iterate_pages(self) -> Generator[Tuple[int, "fitz.Page"], None, None]:
        """
        Yield each page in the PDF using `fitz`.

        Returns:
            Generator[fitz.Page]: Generator yielding each page.
        """
        if fitz is None:
            raise LangroidImportError(
                "pymupdf4llm", ["pymupdf4llm", "all", "pdf-parsers", "doc-chat"]
            )
        doc: fitz.Document = fitz.open(stream=self.doc_bytes, filetype="pdf")
        pages: List[Dict[str, Any]] = pymupdf4llm.to_markdown(doc, page_chunks=True)
        for i, page in enumerate(pages):
            yield i, page
        doc.close()

    def get_document_from_page(self, page: Dict[str, Any]) -> Document:
        """
        Get Document object corresponding to a given "page-chunk"
        dictionary, see:
         https://pymupdf.readthedocs.io/en/latest/pymupdf4llm/api.html


        Args:
            page (Dict[str,Any]): The "page-chunk" dictionary.

        Returns:
            Document: Document object, with content and possible metadata.
        """
        return Document(
            content=self.fix_text(page.get("text", "")),
            # TODO could possible use other metadata from page, see above link.
            metadata=DocMetaData(source=self.source),
        )


class DoclingParser(DocumentParser):
    """
    Parser for processing PDFs using the `docling` library.
    """

    def iterate_pages(self) -> Generator[Tuple[int, Any], None, None]:
        """
        Yield each page in the PDF using `docling`.
        Code largely from this example:
        https://github.com/DS4SD/docling/blob/4d41db3f7abb86c8c65386bf94e7eb0bf22bb82b/docs/examples/export_figures.py

        Returns:
            Generator[docling.Page]: Generator yielding each page.
        """
        if docling is None:
            raise LangroidImportError(
                "docling", ["docling", "pdf-parsers", "all", "doc-chat"]
            )

        from docling.datamodel.base_models import InputFormat  # type: ignore
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        from docling.document_converter import (  # type: ignore
            ConversionResult,
            DocumentConverter,
            PdfFormatOption,
        )
        from docling_core.types.doc import ImageRefMode  # type: ignore

        IMAGE_RESOLUTION_SCALE = 2.0
        pipeline_options = PdfPipelineOptions()
        pipeline_options.images_scale = IMAGE_RESOLUTION_SCALE
        pipeline_options.generate_page_images = True
        pipeline_options.generate_picture_images = True

        converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
            }
        )
        doc_path = self.source
        if doc_path == "bytes":
            # write to tmp file, then use that path
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                temp_file.write(self.doc_bytes.getvalue())
                doc_path = temp_file.name

        output_dir = Path(str(Path(doc_path).with_suffix("")) + "-pages")
        os.makedirs(output_dir, exist_ok=True)

        result: ConversionResult = converter.convert(doc_path)

        def n_page_elements(page) -> int:  # type: ignore
            if page.assembled is None:
                return 0
            return 1 + len(page.assembled.elements)

        page_element_count = [n_page_elements(i) for i in result.pages]
        element_page_cutoff = list(accumulate([1] + page_element_count))
        for i, page in enumerate(result.pages):
            page_start = element_page_cutoff[i]
            page_end = element_page_cutoff[i + 1]
            md_file = output_dir / f"page_{i}.md"
            # we could have just directly exported to a markdown string,
            # but we need to save to a file to force generation of image-files.
            result.document.save_as_markdown(
                md_file,
                image_mode=ImageRefMode.REFERENCED,
                from_element=page_start,
                to_element=page_end,
            )
            yield i, md_file

    def get_document_from_page(self, md_file: str) -> Document:
        """
        Get Document object from a given 1-page markdown file,
        possibly containing image refs.

        Args:
            md_file (str): The markdown file path for the page.

        Returns:
            Document: Document object, with content and possible metadata.
        """
        with open(md_file, "r") as f:
            text = f.read()
        return Document(
            content=self.fix_text(text),
            metadata=DocMetaData(source=self.source),
        )


class PyPDFParser(DocumentParser):
    """
    Parser for processing PDFs using the `pypdf` library.
    """

    def iterate_pages(self) -> Generator[Tuple[int, pypdf.PageObject], None, None]:
        """
        Yield each page in the PDF using `pypdf`.

        Returns:
            Generator[pypdf.pdf.PageObject]: Generator yielding each page.
        """
        if pypdf is None:
            raise LangroidImportError("pypdf", "pdf-parsers")
        reader = pypdf.PdfReader(self.doc_bytes)
        for i, page in enumerate(reader.pages):
            yield i, page

    def get_document_from_page(self, page: pypdf.PageObject) -> Document:
        """
        Get Document object from a given `pypdf` page.

        Args:
            page (pypdf.pdf.PageObject): The `pypdf` page object.

        Returns:
            Document: Document object, with content and possible metadata.
        """
        return Document(
            content=self.fix_text(page.extract_text()),
            metadata=DocMetaData(source=self.source),
        )


class ImagePdfParser(DocumentParser):
    """
    Parser for processing PDFs that are images, i.e. not "true" PDFs.
    """

    def iterate_pages(
        self,
    ) -> Generator[Tuple[int, "Image"], None, None]:  # type: ignore
        try:
            from pdf2image import convert_from_bytes
        except ImportError:
            raise LangroidImportError("pdf2image", "pdf-parsers")

        images = convert_from_bytes(self.doc_bytes.getvalue())
        for i, image in enumerate(images):
            yield i, image

    def get_document_from_page(self, page: "Image") -> Document:  # type: ignore
        """
        Get Document object corresponding to a given `pdf2image` page.

        Args:
            page (Image): The PIL Image object.

        Returns:
            Document: Document object, with content and possible metadata.
        """
        try:
            import pytesseract
        except ImportError:
            raise LangroidImportError("pytesseract", "pdf-parsers")

        text = pytesseract.image_to_string(page)
        return Document(
            content=self.fix_text(text),
            metadata=DocMetaData(source=self.source),
        )


class UnstructuredPDFParser(DocumentParser):
    """
    Parser for processing PDF files using the `unstructured` library.
    """

    def iterate_pages(self) -> Generator[Tuple[int, Any], None, None]:  # type: ignore
        try:
            from unstructured.partition.pdf import partition_pdf
        except ImportError:
            raise ImportError(
                """
                The `unstructured` library is not installed by default with langroid.
                To include this library, please install langroid with the
                `unstructured` extra by running `pip install "langroid[unstructured]"`
                or equivalent.
                """
            )

        # from unstructured.chunking.title import chunk_by_title

        try:
            elements = partition_pdf(file=self.doc_bytes, include_page_breaks=True)
        except Exception as e:
            raise Exception(
                f"""
                Error parsing PDF: {e}
                The `unstructured` library failed to parse the pdf.
                Please try a different library by setting the `library` field
                in the `pdf` section of the `parsing` field in the config file.
                Other supported libraries are:
                fitz, pymupdf4llm, pypdf
                """
            )

        # elements = chunk_by_title(elements)
        page_number = 1
        page_elements = []  # type: ignore
        for el in elements:
            if el.category == "PageBreak":
                if page_elements:  # Avoid yielding empty pages at the start
                    yield page_number, page_elements
                page_number += 1
                page_elements = []
            else:
                page_elements.append(el)
        # Yield the last page if it's not empty
        if page_elements:
            yield page_number, page_elements

    def get_document_from_page(self, page: Any) -> Document:
        """
        Get Document object from a given `unstructured` element.

        Args:
            page (unstructured element): The `unstructured` element object.

        Returns:
            Document: Document object, with content and possible metadata.
        """
        text = " ".join(el.text for el in page)
        return Document(
            content=self.fix_text(text),
            metadata=DocMetaData(source=self.source),
        )


class UnstructuredDocxParser(DocumentParser):
    """
    Parser for processing DOCX files using the `unstructured` library.
    """

    def iterate_pages(self) -> Generator[Tuple[int, Any], None, None]:  # type: ignore
        try:
            from unstructured.partition.docx import partition_docx
        except ImportError:
            raise ImportError(
                """
                The `unstructured` library is not installed by default with langroid.
                To include this library, please install langroid with the
                `unstructured` extra by running `pip install "langroid[unstructured]"`
                or equivalent.
                """
            )

        elements = partition_docx(file=self.doc_bytes, include_page_breaks=True)

        page_number = 1
        page_elements = []  # type: ignore
        for el in elements:
            if el.category == "PageBreak":
                if page_elements:  # Avoid yielding empty pages at the start
                    yield page_number, page_elements
                page_number += 1
                page_elements = []
            else:
                page_elements.append(el)
        # Yield the last page if it's not empty
        if page_elements:
            yield page_number, page_elements

    def get_document_from_page(self, page: Any) -> Document:
        """
        Get Document object from a given `unstructured` element.

        Note:
            The concept of "pages" doesn't actually exist in the .docx file format in
            the same way it does in formats like .pdf. A .docx file is made up of a
            series of elements like paragraphs and tables, but the division into
            pages is done dynamically based on the rendering settings (like the page
            size, margin size, font size, etc.).

        Args:
            page (unstructured element): The `unstructured` element object.

        Returns:
            Document object, with content and possible metadata.
        """
        text = " ".join(el.text for el in page)
        return Document(
            content=self.fix_text(text),
            metadata=DocMetaData(source=self.source),
        )


class UnstructuredDocParser(UnstructuredDocxParser):
    def iterate_pages(self) -> Generator[Tuple[int, Any], None, None]:  # type: ignore
        try:
            from unstructured.partition.doc import partition_doc
        except ImportError:
            raise ImportError(
                """
                The `unstructured` library is not installed by default with langroid.
                To include this library, please install langroid with the
                `unstructured` extra by running `pip install "langroid[unstructured]"`
                or equivalent.
                """
            )

        elements = partition_doc(file=self.doc_bytes, include_page_breaks=True)

        page_number = 1
        page_elements = []  # type: ignore
        for el in elements:
            if el.category == "PageBreak":
                if page_elements:  # Avoid yielding empty pages at the start
                    yield page_number, page_elements
                page_number += 1
                page_elements = []
            else:
                page_elements.append(el)
        # Yield the last page if it's not empty
        if page_elements:
            yield page_number, page_elements


class PythonDocxParser(DocumentParser):
    """
    Parser for processing DOCX files using the `python-docx` library.
    """

    def iterate_pages(self) -> Generator[Tuple[int, Any], None, None]:
        """
        Simulate iterating through pages.
        In a DOCX file, pages are not explicitly defined,
        so we consider each paragraph as a separate 'page' for simplicity.
        """
        try:
            import docx
        except ImportError:
            raise LangroidImportError("python-docx", "docx")

        doc = docx.Document(self.doc_bytes)
        for i, para in enumerate(doc.paragraphs, start=1):
            yield i, [para]

    def get_document_from_page(self, page: Any) -> Document:
        """
        Get Document object from a given 'page', which in this case is a single
        paragraph.

        Args:
            page (list): A list containing a single Paragraph object.

        Returns:
            Document: Document object, with content and possible metadata.
        """
        paragraph = page[0]
        return Document(
            content=self.fix_text(paragraph.text),
            metadata=DocMetaData(source=self.source),
        )


class MarkitdownXLSXParser(DocumentParser):
    def iterate_pages(self) -> Generator[Tuple[int, Any], None, None]:
        try:
            from markitdown import MarkItDown
        except ImportError:
            LangroidImportError("markitdown", "doc-parsers")
        md = MarkItDown()
        self.doc_bytes.seek(0)  # Reset to start

        # Save stream to a temp file since md.convert() expects a path or URL
        # Temporary workaround until markitdown fixes convert_stream function
        # for xls and xlsx files
        # See issue here https://github.com/microsoft/markitdown/issues/321
        with tempfile.NamedTemporaryFile(delete=True, suffix=".xlsx") as temp_file:
            temp_file.write(self.doc_bytes.read())
            temp_file.flush()  # Ensure data is written before reading
            result = md.convert(temp_file.name)

        sheets = re.split(r"(?=## Sheet\d+)", result.text_content)

        for i, sheet in enumerate(sheets):
            yield i, sheet

    def get_document_from_page(self, md_content: str) -> Document:
        """
        Get Document object from a given 1-page markdown string.

        Args:
            md_content (str): The markdown content for the page.

        Returns:
            Document: Document object, with content and possible metadata.
        """
        return Document(
            content=self.fix_text(md_content),
            metadata=DocMetaData(source=self.source),
        )


class MarkitdownPPTXParser(DocumentParser):
    def iterate_pages(self) -> Generator[Tuple[int, Any], None, None]:
        try:
            from markitdown import MarkItDown
        except ImportError:
            LangroidImportError("markitdown", "doc-parsers")

        md = MarkItDown()
        self.doc_bytes.seek(0)
        result = md.convert_stream(self.doc_bytes, file_extension=".pptx")
        slides = re.split(r"(?=<!-- Slide number: \d+ -->)", result.text_content)
        for i, slide in enumerate(slides):
            yield i, slide

    def get_document_from_page(self, md_content: str) -> Document:
        """
        Get Document object from a given 1-page markdown string.

        Args:
            md_content (str): The markdown content for the page.

        Returns:
            Document: Document object, with content and possible metadata.
        """
        return Document(
            content=self.fix_text(md_content),
            metadata=DocMetaData(source=self.source),
        )
